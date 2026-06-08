#!/usr/bin/env python3
"""Generate simple superior questions docx."""

from docx import Document
from docx.shared import Pt

QUESTIONS = [
    ("Geography & people", [
        "What are the official pilot cluster (village) names?",
        "Who are the Resham Doots — name, phone, email — and which cluster is each assigned to?",
        "Who are the central admin(s) — name, phone, email?",
        "Can one Resham Doot cover more than one cluster?",
    ]),
    ("Categories & skills", [
        "What is the full producer category list? (It is hierarchical and skill-level based.)",
        "Can one producer have multiple categories? If yes, is there a primary category?",
        "What is the full skill tags list?",
        "Are skills predefined only, or can field staff add free text?",
    ]),
    ("Income & required fields", [
        "What are the exact income range bands (rupee brackets)?",
        "Which fields are mandatory at onboarding besides name, photo, phone, date of birth, cluster, and skills?",
        "Should we collect alternate / family phone?",
        "Should we collect household size and number of dependents?",
        "Should we collect SHG / cooperative name?",
        "Should we collect machine / tool access?",
    ]),
    ("Government schemes", [
        "Which 10–15 schemes should be in the pilot list?",
        "What eligibility rules apply for each (gender, category, age — no BPL)?",
        "What disclaimer text should appear on matched schemes?",
        "Who owns updating the scheme list and how often?",
    ]),
    ("Producer card (digital)", [
        "Which fields appear on the digital producer card?",
        "Please approve card visual design (logo, colours, layout). Physical printing is deferred.",
    ]),
    ("Pilot plan", [
        "What is the pilot start date and timeline?",
        "What are success criteria? (e.g. number of producers with photo and card)",
        "Which funders or government partners get the read-only dashboard link?",
        "How often should reports be shared?",
    ]),
    ("Language", [
        "Confirm languages for MVP v1: English, Hindi, and Assamese all at once, or phased? (Demo already ships EN + HI + AS for core UI.)",
        "Should translations be human-reviewed or machine-translated for dynamic content (scheme names, skill labels, govt copy)?",
    ]),
    ("Dashboard", [
        "Which 8–10 metrics does leadership want in monthly pilot reviews?",
    ]),
]

def main():
    doc = Document()
    title = doc.add_heading("Producer System — Questions for Leadership", level=0)
    title.runs[0].font.size = Pt(16)
    doc.add_paragraph(
        "Please fill in answers below. These are the only items blocking MVP content and rules. "
        "Locked product decisions and answered build questions are documented in questions.md."
    )
    doc.add_paragraph("")

    n = 1
    for section, items in QUESTIONS:
        doc.add_heading(section, level=1)
        for q in items:
            p = doc.add_paragraph()
            p.add_run(f"{n}. {q}").bold = False
            doc.add_paragraph("Answer: _______________________________________________")
            doc.add_paragraph("")
            n += 1

    out = "Producer_card/questions-for-superior.docx"
    import os
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    path = os.path.join(base, "questions-for-superior.docx")
    doc.save(path)
    print(f"Wrote {path}")

if __name__ == "__main__":
    main()
